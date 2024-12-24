import math
from openpyxl import Workbook, load_workbook
from docx import Document
import sqlite3
import os


def vvod(var):
    while True:
        try:
            return float(input(f'{var}: ').replace(',', '.'))
        except:
            print('Введите число!')


def vivod(var, res):
    print(f'Переменная {var} = {res}')


def solveS(x):
    s = 1 + x + math.pow(x, 2)/2 + math.pow(x, 3) / 6 + math.pow(x, 4) / 24
    return s


def solveF(sf, x, y):
    sf['f'] = x * (math.sin(math.pow(x, 3)) + math.pow(math.cos(y), 2))


def solveSF(sf, x, y):
    sf['s'] = solveS(x)
    solveF(sf, x, y)


def save_to_file(filename, text, value):
    """
    Сохраняет данные в файл.
    :param filename: имя файла
    :param text: подпись над данными
    :param value: данные
    """
    with open(filename, 'a', encoding='utf-8') as file:
        file.write(f"{text}:\n")
        if isinstance(value, list):
            file.write(", ".join(map(str, value)) + "\n\n")
        else:
            file.write(f"{value}\n\n")


def save_to_excel(filename, text, value):
    """
    Записывает данные в excel
    :param filename: путь к excel таблице
    :param text: текст над данными
    :param value: данные
    """
    if os.path.exists(filename):
        wb = load_workbook(filename)
    else:
        wb = Workbook()

    ws = wb.active
    ws.append([])

    ws.append([text])
    if isinstance(value, list):
        ws.append(value)
    else:
        ws.append([value])
    wb.save(filename)


def save_to_word(filename, text, value):
    """
    Создаёт word файл с таблицей
    :param filename: путь к word документу
    :param text: текст над данными
    :param value: данные
    """
    if os.path.exists(filename):
        doc = Document(filename)
    else:
        doc = Document()

    doc.add_heading(f'{text}:', level=1)

    if isinstance(value, list):
        doc.add_paragraph(', '.join(map(str, value)))
    else:
        doc.add_paragraph(str(value))
    doc.save(filename)


def connect_sqlite(filename):
    """
    подкючение к SQLite таблице
    :param filename: имя файла
    :return: объект коннектора
    """
    connector = sqlite3.connect(filename)
    connector.cursor().execute("CREATE TABLE IF NOT EXISTS results (id INTEGER PRIMARY KEY AUTOINCREMENT)")
    return connector


def save_to_sqlite(filename, name, value):
    """
    Сохраняет данные в таблицу SQLite
    :param filename: имя файла
    :param name: название столбца
    :param value: данные
    """
    connector = connect_sqlite(filename)
    cursor = connector.cursor()

    cursor.execute(f"PRAGMA table_info(results)")
    columns = [row[1] for row in cursor.fetchall()]
    if name not in columns:
        cursor.execute(f"ALTER TABLE results ADD COLUMN {name} TEXT")

    cursor.execute("SELECT COUNT(*) FROM results")
    row_count = cursor.fetchone()[0]

    if isinstance(value, list):
        for i in range(min(row_count, len(value))):
            cursor.execute(f"UPDATE results SET {name} = {value[i]} WHERE rowid = {i + 1}")

        for i in range(row_count, len(value)):
            cursor.execute(f"INSERT INTO results ({name}) VALUES ({value[i]})")
    else:
        if row_count:
            cursor.execute(f"UPDATE results SET {name} = {value} WHERE rowid = 1")
        else:
            cursor.execute(f"INSERT INTO results ({name}) VALUES ({value})")

    connector.commit()
    connector.close()


def save2(arr, res, resarr):
    save_to_word(f'DBsaves\PyWord2Lab.docx', 'Сгенерированный массив', arr)
    save_to_word(f'DBsaves\PyWord2Lab.docx', 'Индекс максимального четного элемента', res)
    save_to_word(f'DBsaves\PyWord2Lab.docx', 'Новый массив', resarr)
    save_to_excel(f'DBsaves\PyExcel2Lab.xlsx', 'Сгенерированный массив', arr)
    save_to_excel(f'DBsaves\PyExcel2Lab.xlsx', 'Индекс максимального четного элемента', res)
    save_to_excel(f'DBsaves\PyExcel2Lab.xlsx', 'Новый массив', resarr)
    save_to_sqlite(f'DBsaves\PySQLite2Lab.db', 'Generated_array', arr)
    save_to_sqlite(f'DBsaves\PySQLite2Lab.db', 'Index_of_max_chet_value', res)
    save_to_sqlite(f'DBsaves\PySQLite2Lab.db', 'Result_array', resarr)
    save_to_file(f'DBsaves\PyTextFile2Lab.txt', 'Сгенерированный массив', arr)
    save_to_file(f'DBsaves\PyTextFile2Lab.txt', 'Индекс максимального четного элемента', res)
    save_to_file(f'DBsaves\PyTextFile2Lab.txt', 'Новый массив', resarr)


def save3(genarr, results):
    tasks = [3, 4, 5, 6, 7, 8, 9, 10, 12]
    save_to_word(f'DBsaves\PyWord3Lab.docx', f'Сгенерированный массив', genarr)
    save_to_excel(f'DBsaves\PyExcel3Lab.xlsx', f'Сгенерированный массив', genarr)
    save_to_sqlite(f'DBsaves\PySQLite3Lab.db', f'Generated', genarr)
    save_to_file(f'DBsaves\PyTextFile3Lab.txt', f'Сгенерированный массив', genarr)

    for i in range(9):
        save_to_word(f'DBsaves\PyWord3Lab.docx', f'Ответ {tasks[i]}', results[i])
        save_to_excel(f'DBsaves\PyExcel3Lab.xlsx', f'Ответ {tasks[i]}', results[i])
        save_to_sqlite(f'DBsaves\PySQLite3Lab.db', f'Answer_{tasks[i]}', results[i])
        save_to_file(f'DBsaves\PyTextFile3Lab.txt', f'Ответ {tasks[i]}', results[i])
